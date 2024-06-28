import pandas as pd
import numpy as np 
import seaborn as sb
import matplotlib.pyplot as plt 
import subprocess
from docx import Document
from docx.shared import Inches


def train_data_analysis():
    
    doc = Document()
    
    result = subprocess.run(['python', 'data_analysis.py'],capture_output=True, text=True)
    if result.returncode!=0:
        print(f"Error running data_analysis.py: {result.stderr}")
    else:
        print(f"Successfully ran data_analysis.py")
        #print(result.stdout)
        
           
        #### Saving Data analyses to a Document 
        doc.add_heading('Data Analysis Report', 0)
        doc.add_heading('Summary Statistics', level=1)
        
        # Add the description text directly from stdout
        doc.add_paragraph(result.stdout)
        
        # Add the plots images
        doc.add_heading('Loan Status Plot', level=1)
        doc.add_picture('loan_status_plot.png', width=Inches(6.0))
        
        doc.add_heading('Plot of Categorical Variables', level=1)
        doc.add_picture('plots_categorical.png', width=Inches(6.0))
        
        doc.add_heading('Plot of Ordinal Variables', level=1)
        doc.add_picture('plots_ordinals.png', width=Inches(6.0))
        
        
        doc.save('data_analysis_report.docx')
    

def train():

    print("")


    









if __name__== '__main__':

    ## converting training data to a dataframe
    #train_df = pd.read_csv('train.csv')
    train_data_analysis()
    #train(train_df)

    